from __future__ import annotations

from collections import defaultdict
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("全国小学课外书权威扩展书单.docx")

ACCENT = "0F4C5C"
ACCENT_2 = "DDAA33"
INK = "172A3A"
MUTED = "667085"
FILL = "EAF3F5"
FILL_2 = "F7FBFC"
HEADER_FILL = "0F4C5C"
WHITE = "FFFFFF"
BORDER = "D0D5DD"


def rgb(hex_color: str) -> RGBColor:
    h = hex_color.strip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_run_font(run, size=None, color=None, bold=None, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text, bold=False, color=INK, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.16
    r = p.add_run(str(text))
    set_run_font(r, size=size, color=color, bold=bold, font="Microsoft YaHei")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_run_font(r, size=17.2 if level == 1 else 13.4, color=ACCENT if level == 1 else INK, bold=True)
    return p


def add_body(doc, text, color=INK, size=11.4, after=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.24
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, font="Microsoft YaHei")
    return p


def add_source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=9.3, color=MUTED, font="Microsoft YaHei")


def add_table(doc, headers, rows, widths, header_fill=HEADER_FILL, zebra=True, font_size=9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    set_table_widths(table, widths)
    hdr = table.rows[0].cells
    repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        shade_cell(hdr[i], header_fill)
        set_cell_text(hdr[i], h, bold=True, color=WHITE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx, value in enumerate(row):
            if zebra and ridx % 2:
                shade_cell(cells[cidx], FILL_2)
            set_cell_text(
                cells[cidx],
                value,
                size=font_size,
                align=WD_ALIGN_PARAGRAPH.CENTER if cidx in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT,
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def parse_guidance_books():
    md_path = Path("全国小学课外阅读政策依据与权威书单.md")
    if not md_path.exists():
        raise FileNotFoundError("缺少上一版已校验的 Markdown 数据源")

    books = []
    stage = None
    for line in md_path.read_text(encoding="utf-8").splitlines():
        if line.startswith("### 小学") and "年级" in line:
            stage = line.replace("### ", "").strip()
            continue
        if not stage or not line.startswith("|"):
            continue
        if line.startswith("|---") or "图书名称" in line or "序号" in line:
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if len(cells) == 4 and cells[0].isdigit():
            no = int(cells[0])
            if 1 <= no <= 110:
                books.append((no, stage, cells[1], cells[2], cells[3]))

    if len(books) != 110:
        raise RuntimeError(f"阅读指导目录解析得到 {len(books)} 种，预期 110 种")
    return books


CURRICULUM_EXAMPLES = [
    ("课标附录举例", "童话", "安徒生童话", "安徒生"),
    ("课标附录举例", "童话", "格林童话", "格林兄弟"),
    ("课标附录举例", "童话", "稻草人", "叶圣陶"),
    ("课标附录举例", "童话", "宝葫芦的秘密", "张天翼"),
    ("课标附录举例", "寓言", "中国古今寓言", "多作者"),
    ("课标附录举例", "寓言", "伊索寓言", "伊索"),
    ("课标附录举例", "故事", "成语故事", "多作者"),
    ("课标附录举例", "故事", "神话故事", "多作者"),
    ("课标附录举例", "故事", "民间故事", "多作者"),
    ("课标附录举例", "故事", "中外历史故事", "多作者"),
    ("课标附录举例", "诗歌散文", "朝花夕拾", "鲁迅"),
    ("课标附录举例", "诗歌散文", "繁星·春水", "冰心"),
    ("课标附录举例", "诗歌散文", "艾青诗选", "艾青"),
    ("课标附录举例", "诗歌散文", "可爱的中国", "方志敏"),
    ("课标附录举例", "诗歌散文", "革命烈士诗抄", "萧三等编"),
    ("课标附录举例", "长篇名著", "西游记", "吴承恩"),
    ("课标附录举例", "长篇名著", "水浒传", "施耐庵"),
    ("课标附录举例", "长篇名著", "骆驼祥子", "老舍"),
    ("课标附录举例", "长篇名著", "红岩", "罗广斌、杨益言"),
    ("课标附录举例", "长篇名著", "红星照耀中国", "埃德加·斯诺"),
    ("课标附录举例", "长篇名著", "格列佛游记", "斯威夫特"),
    ("课标附录举例", "长篇名著", "简·爱", "夏洛蒂·勃朗特"),
    ("课标附录举例", "长篇名著", "童年", "高尔基"),
    ("课标附录举例", "长篇名著", "钢铁是怎样炼成的", "奥斯特洛夫斯基"),
    ("课标附录举例", "科普科幻", "十万个为什么", "多作者/版本"),
    ("课标附录举例", "科普科幻", "海底两万里", "儒勒·凡尔纳"),
]


HAPPY_READING = [
    ("一年级上册", "和大人一起读（全四册）", "张秋生", "二十一世纪出版社"),
    ("一年级下册", "读读童谣和儿歌（全四册）", "李宏声", "二十一世纪出版社"),
    ("二年级上册", "小鲤鱼跳龙门", "金近", "吉林出版集团"),
    ("二年级上册", "小狗的小房子", "孙幼军", "吉林出版集团"),
    ("二年级上册", "一只想飞的猫", "陈伯吹", "吉林出版集团"),
    ("二年级上册", "“歪脑袋”木头桩", "严文井", "吉林出版集团"),
    ("二年级上册", "孤独的小螃蟹", "冰波", "青岛出版社"),
    ("二年级下册", "一起长大的玩具－怪手杖", "金波", "青岛出版社"),
    ("二年级下册", "神笔马良", "洪汛涛", "北京燕山出版社"),
    ("二年级下册", "七色花", "瓦·卡达耶夫/著，曹靖华/译", "北京燕山出版社"),
    ("二年级下册", "愿望的实现", "泰戈尔/著，郑振铎/译", "北京燕山出版社"),
    ("二年级下册", "大头儿子和小头爸爸", "央视动漫集团", "北京燕山出版社"),
    ("三年级上册", "格林童话", "格林兄弟/著，杨武能/译", "二十一世纪出版社"),
    ("三年级上册", "安徒生童话", "安徒生/著，叶君健/译", "二十一世纪出版社"),
    ("三年级上册", "稻草人", "叶圣陶", "二十一世纪出版社"),
    ("三年级下册", "中国古代寓言", "溪石", "吉林出版集团"),
    ("三年级下册", "伊索寓言", "伊索/著，杨立新/译", "吉林出版集团"),
    ("三年级下册", "克雷洛夫寓言", "克雷洛夫/著，韦苇/译", "吉林出版集团"),
    ("四年级上册", "中国古代神话", "乌兰托娅", "吉林出版集团"),
    ("四年级上册", "世界神话传说", "卢艳丽", "吉林出版集团"),
    ("四年级上册", "古希腊神话", "版本信息未列作者", "吉林出版集团"),
    ("四年级下册", "十万个为什么", "米·伊林", "北京燕山出版社"),
    ("四年级下册", "森林报", "维·比安基", "北京燕山出版社"),
    ("四年级下册", "细菌世界历险记", "高士其", "北京燕山出版社"),
    ("四年级下册", "地球的故事", "房龙", "北京燕山出版社"),
    ("四年级下册", "爷爷的爷爷哪里来", "贾兰坡", "北京燕山出版社"),
    ("四年级下册", "穿过地平线", "李四光", "北京燕山出版社"),
    ("五年级上册", "中国民间故事", "卢丽艳/编", "吉林出版集团"),
    ("五年级上册", "非洲民间故事", "版本信息未列作者", "吉林出版集团"),
    ("五年级上册", "欧洲民间故事", "版本信息未列作者", "吉林出版集团"),
    ("五年级上册", "列那狐的故事", "[法]吉罗夫人/著，张国魁/编", "吉林出版集团"),
    ("五年级下册", "西游记", "吴承恩", "北京燕山出版社"),
    ("五年级下册", "三国演义", "罗贯中", "北京燕山出版社"),
    ("五年级下册", "水浒传", "施耐庵", "北京燕山出版社"),
    ("五年级下册", "红楼梦", "曹雪芹", "北京燕山出版社"),
    ("六年级上册", "童年", "[苏]高尔基/著，高晓慧/译", "吉林出版集团"),
    ("六年级上册", "小英雄雨来", "管桦", "吉林出版集团"),
    ("六年级上册", "爱的教育", "[意]亚米契斯/著，高晓慧/译", "吉林出版集团"),
    ("六年级下册", "鲁滨逊漂流记", "丹尼尔·笛福", "北京燕山出版社"),
    ("六年级下册", "骑鹅旅行记", "塞尔玛·拉格洛夫", "北京燕山出版社"),
    ("六年级下册", "爱丽丝漫游奇境", "刘易斯·卡罗尔", "北京燕山出版社"),
    ("六年级下册", "汤姆·索亚历险记", "马克·吐温", "辽海出版社"),
]


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    for style_name in ("Normal", "Body Text"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11.4)
        style.font.color.rgb = rgb(INK)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("全国小学课外书权威扩展书单")
    set_run_font(run, size=9.2, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("依据国家课标、教育部阅读指导目录与统编语文教材栏目整理")
    set_run_font(r, size=8.8, color=MUTED)


def add_cover(doc, guidance_count):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("小学课外书权威扩展书单")
    set_run_font(r, size=28, color=ACCENT, bold=True)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("1-6年级｜课标示例 × 教材栏目 × 教育部指导目录")
    set_run_font(r, size=14.2, color=MUTED)
    p.paragraph_format.space_after = Pt(24)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    set_table_widths(table, [2.2, 2.5, 2.2])
    set_table_borders(table, color="FFFFFF")
    stats = [("26", "课标附录点名示例"), (str(len(HAPPY_READING)), "语文教材栏目关联书目"), (str(guidance_count), "教育部阅读指导目录小学段")]
    for idx, (num, label) in enumerate(stats):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, FILL)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(num)
        set_run_font(r, size=24, color=ACCENT, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, size=9.2, color=INK)

    doc.add_paragraph()
    add_body(
        doc,
        "说明：本文件只收录课外书/课外读物，不收录古诗词背诵篇目。国家课标没有发布逐年级固定必读书单，本文按来源分层列出可公开核验的权威范围。",
        color=INK,
        size=11.4,
        after=12,
    )


def grouped_happy_rows():
    return [(grade, title, author, publisher) for grade, title, author, publisher in HAPPY_READING]


def main():
    guidance_books = parse_guidance_books()

    doc = Document()
    configure_doc(doc)
    add_cover(doc, len(guidance_books))

    add_heading(doc, "一、使用口径", 1)
    add_body(doc, "本书单分三类来源：第一类是《义务教育语文课程标准（2022年版）》附录中举例出现的课外读物；第二类是统编小学语文教科书“快乐读书吧”栏目关联书目；第三类是教育部基础教育课程教材发展中心《中小学生阅读指导目录（2020年版）》小学段110种。")
    add_body(doc, "三类来源的法律/政策强度不同，因此文档中保留来源区分。重复书名不合并，便于判断同一本书在不同政策或教材场景中的出现位置。", color=MUTED, size=10.4)

    add_heading(doc, "二、课标附录点名的课外读物示例", 1)
    add_source_note(doc, "来源：《义务教育语文课程标准（2022年版）》附录“关于课内外读物的建议”。这些是课标举例，不是小学逐年级必读目录。")
    add_table(
        doc,
        ["来源", "类型", "书名/读物", "作者/说明"],
        CURRICULUM_EXAMPLES,
        [1.3, 1.0, 2.65, 1.95],
        font_size=9.1,
    )

    doc.add_page_break()
    add_heading(doc, "三、统编小学语文教材“快乐读书吧”关联书目", 1)
    add_source_note(doc, "来源：四川省教育厅2022年“春日阳光阅读”附件1列出的“统编版小学语文教科书的特色栏目‘快乐读书吧’”版本信息，并参照人民教育出版社“快乐读书吧”栏目资料校验。版本信息用于定位阅读内容，不代表只能使用该出版社版本。")
    add_table(
        doc,
        ["册次", "书名", "作者/译者/编者", "版本信息"],
        grouped_happy_rows(),
        [1.05, 2.25, 2.3, 1.3],
        font_size=8.7,
    )

    doc.add_page_break()
    add_heading(doc, "四、教育部阅读指导目录小学段110种", 1)
    add_source_note(doc, "来源：教育部基础教育课程教材发展中心《中小学生阅读指导目录（2020年版）》小学段。该目录供学生自主选择阅读，各地各校不作统一要求。")
    by_stage = defaultdict(list)
    for no, stage, cat, title, author in guidance_books:
        by_stage[stage].append((no, cat, title, author))
    for stage in ("小学1-2年级", "小学3-4年级", "小学5-6年级"):
        add_heading(doc, stage, 2)
        add_table(
            doc,
            ["序号", "类别", "书名", "作者/编者"],
            by_stage[stage],
            [0.55, 0.85, 2.85, 2.65],
            font_size=8.5,
        )

    doc.add_page_break()
    add_heading(doc, "五、官方来源与引用边界", 1)
    sources = [
        ("义务教育课程方案和课程标准（2022年版）", "教育部通知", "https://www.moe.gov.cn/srcsite/A26/s8001/202204/t20220420_619921.html"),
        ("义务教育语文课程标准（2022年版）", "教育部制定", "https://jyxy.xawl.edu.cn/__local/2/B6/04/716DD93B3176A55F7D541D0E842_8C9345E6_13B7630.pdf?e=.pdf"),
        ("中小学生阅读指导目录（2020年版）", "教育部基础教育课程教材发展中心", "https://www.moe.gov.cn/jyb_xwfb/gzdt_gzdt/s5987/202004/W020200422556593462993.pdf"),
        ("中小学生课外读物进校园管理办法", "教育部，教材〔2021〕2号", "https://www.moe.gov.cn/srcsite/A26/moe_714/202104/t20210401_523904.html"),
        ("快乐读书吧", "人民教育出版社品牌丛书页面", "https://www.pep.com.cn/products/tj/ppcs/jf/kldsb/"),
        ("春日阳光阅读推荐书目版本信息", "四川省教育厅附件", "https://www.zhongjiang.gov.cn/wcm.files/upload/202206/202206170428026.pdf"),
    ]
    add_table(doc, ["资料", "来源/性质", "链接"], sources, [2.0, 1.5, 3.4], font_size=8.4)
    add_body(doc, "边界说明：本文件不使用“教育部推荐”“新课标指定”等商业化表述；凡属课标举例、教材栏目关联、阅读指导目录，均分开标注，避免把不同效力的来源混同为全国强制必读。", color=MUTED, size=10.2)

    doc.save(OUT)
    print(OUT.resolve())
    print(f"happy_reading={len(HAPPY_READING)} guidance={len(guidance_books)} curriculum_examples={len(CURRICULUM_EXAMPLES)}")


if __name__ == "__main__":
    main()
